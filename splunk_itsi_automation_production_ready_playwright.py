# Project: Splunk ITSI Multi-Portal Automation (Production Ready)
# Author: Atul Solanki
# Description:
# Automates login (SSO + DUO manual once), runs queries, captures screenshots,
# and generates a Word report across multiple ITSI portals.

# =========================
# 1. Setup Instructions
# =========================
# python -m venv venv
# source venv/bin/activate   (Linux/Mac)
# venv\Scripts\activate      (Windows)
# source venv/Scripts/activate (Git Bash on windows)
# python.exe -m pip install --upgrade pip
# pip install -r requirements.txt
# pip install --upgrade pip setuptools wheel
# pip install greenlet
# pip install playwright python-docx python-dotenv
# playwright install


# =========================
# 2. requirements.txt
# =========================
# playwright==1.42.0
# python-docx==1.1.0
# python-dotenv==1.0.1


# =========================
# 3. main.py
# =========================

import json
import os
from datetime import datetime
from urllib.parse import quote, urlencode, urlsplit, urlunsplit
from playwright.sync_api import TimeoutError as PlaywrightTimeoutError
from playwright.sync_api import Error as PlaywrightError
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dotenv import load_dotenv

load_dotenv()

PORTALS = [
    {"name": "Allianz ACI NG ITSI", "url": "https://cms-itsi-portal.cisco.com/edc0941d/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Amgen NG ITSI", "url": "https://cms-itsi-portal.cisco.com/96264655/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Amex GBT NG ITSI", "url": "https://cms-itsi-portal.cisco.com/974ed1d6/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "CaaS NG ITSI", "url": "https://cms-itsi-portal.cisco.com/39633d56/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Cust0-SNOW NG ITSI", "url": "https://cms-itsi-portal.cisco.com/9a120b19/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Dawiyat NG ITSI", "url": "https://cms-itsi-portal.cisco.com/ccde3249/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Engel NG ITSI", "url": "https://cms-itsi-portal.cisco.com/76393f81/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Ford NGC NG ITSI", "url": "https://cms-itsi-portal.cisco.com/a228b05f/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Gilead NG ITSI", "url": "https://cms-itsi-portal.cisco.com/2096b955/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Global Foundries NG ITSI", "url": "https://cms-itsi-portal.cisco.com/62bb01cc/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Hilton NG ITSI", "url": "https://cms-itsi-portal.cisco.com/2b4576e0/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Humana NG ITSI", "url": "https://cms-itsi-portal.cisco.com/cfcbc475/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Liberty Mutual NG ITSI", "url": "https://cms-itsi-portal.cisco.com/683886c3/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "MGB NG ITSI", "url": "https://cms-itsi-portal.cisco.com/3c1cc9ae/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "MGCS NG ITSI", "url": "https://cms-itsi-portal.cisco.com/4a7cee11/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Michigan Medicine NG ITSI", "url": "https://cms-itsi-portal.cisco.com/d9ee7a35/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Mobily NG ITSI", "url": "https://cms-itsi-portal.cisco.com/1f28e2fd/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "SAP Email Security NG ITSI", "url": "https://cms-itsi-portal.cisco.com/6e28326c/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Societe Generale NG ITSI", "url": "https://cms-itsi-portal.cisco.com/4a07135/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Truist NG ITSI", "url": "https://cms-itsi-portal.cisco.com/1ec025ee/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "UPMC NG ITSI", "url": "https://cms-itsi-portal.cisco.com/8cf530d/en-GB/app/cms_sapphire_gnoc/search"},
    {"name": "Vodafone SDWAN NG ITSI", "url": "https://cms-itsi-portal.cisco.com/d4de824b/en-GB/app/cms_sapphire_gnoc/search"}
]

QUERIES = [
    {
        "name": "Interface Performance",
        "query": "index=snmp_annotated | search command IN (_IF-MIB*, _RFC1213-MIB*, _SNMPv2-MIB*)",
        "earliest": "-15m",
    },
    {
        "name": "Device Performance",
        "query": "index=snmp_annotated command=*_perfData | search NOT command IN (_IF-MIB*, _RFC1213-MIB*, _SNMPv2-MIB*)",
        "earliest": "-15m",
    },
    {
        "name": "Ping Monitor",
        "query": "index=ping_monitor",
        "earliest": "-15m",
    },
    {
        "name": "Group Alert",
        "query": "index=itsi_grouped_alerts",
        "earliest": "-60m",
    }
]

QUERY_OPTION_MAP = {
    "2": "Ping Monitor",
    "3": "Interface Performance",
    "4": "Device Performance",
    "5": "Group Alert",
}
DEFAULT_QUERY_NAMES = ["Interface Performance", "Device Performance", "Ping Monitor"]

OUTPUT_DIR = "output"
SCREENSHOT_DIR = os.path.join(OUTPUT_DIR, "screenshots")
AUTH_STATE_FILE = "auth.json"
NAVIGATION_TIMEOUT_MS = 120000
MANUAL_LOGIN_TIMEOUT_MS = 120000
PORTAL_RETRY_COUNT = 2
SEARCH_RESULTS_WAIT_MS = int(os.getenv("SEARCH_RESULTS_WAIT_MS", "5000"))
SEARCH_TEXTAREA_SELECTOR = "pre.ace_editor.ace-spl-light textarea.ace_text-input"
SEARCH_EDITOR_SELECTOR = "pre.ace_editor.ace-spl-light, .ace_editor.ace-spl-light"
SEARCH_INPUT_SELECTOR = (
    "textarea[placeholder*='enter search here'], "
    "input[placeholder*='enter search here'], "
    "textarea[aria-label='Search query'], "
    "input[aria-label='Search query'], "
    "textarea[data-test='search-field'], "
    "input[data-test='search-field'], "
    "textarea.search-input, "
    "input.search-input"
)
SEARCH_BUTTON_SELECTOR = (
    "a.btn[aria-label='Search Button'], "
    "button[aria-label='Search Button'], "
    "button[aria-label='Search'], "
    "a[aria-label='Search'], "
    "button[title='Search'], "
    "a[title='Search'], "
    "button[data-test='search-button'], "
    "a[data-test='search-button']"
)
TIME_PICKER_BUTTON_SELECTOR = (
    "div.shared-timerangepicker a.btn[aria-haspopup='true'], "
    "div.shared-timerangepicker button[aria-haspopup='true'], "
    "button:has-text('Time range:'), "
    "a:has-text('Time range:')"
)
TIME_PICKER_LABEL_SELECTOR = "div.shared-timerangepicker span.time-label"
TIME_PICKER_DIALOG_SELECTOR = "div[data-render-time].shared-timerangepicker-dialog"
TIME_PICKER_PRESET_15_MINUTES_SELECTOR = (
    f"{TIME_PICKER_DIALOG_SELECTOR} a[data-earliest='-15m'][data-latest='now']"
)
AUTH_REDIRECT_HOSTS = ("id.cisco.com", "duosecurity.com")
os.makedirs(SCREENSHOT_DIR, exist_ok=True)


def create_document():
    doc = Document()
    doc.add_heading('Splunk ITSI Automation Report', 0)
    doc.add_paragraph(f"Generated on: {datetime.now()}")
    return doc


def add_hyperlink(paragraph, url, text):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def login_once(page):
    print(
        "Login required. Please complete SSO + DUO manually "
        f"within {MANUAL_LOGIN_TIMEOUT_MS // 1000} seconds..."
    )
    page.wait_for_timeout(MANUAL_LOGIN_TIMEOUT_MS)


def get_today_string():
    return datetime.now().date().isoformat()


def load_saved_auth_state(auth_state_path):
    if not os.path.exists(auth_state_path):
        return None

    with open(auth_state_path, "r", encoding="utf-8") as auth_file:
        auth_data = json.load(auth_file)

    if "storage_state" in auth_data and "saved_on" in auth_data:
        return auth_data

    # Backward compatibility for older plain Playwright storage state files.
    return {
        "saved_on": datetime.fromtimestamp(os.path.getmtime(auth_state_path)).date().isoformat(),
        "storage_state": auth_data,
    }


def should_refresh_auth_state(auth_state_path):
    saved_auth_state = load_saved_auth_state(auth_state_path)
    if not saved_auth_state:
        return True

    return saved_auth_state["saved_on"] != get_today_string()


def save_auth_state(context, auth_state_path):
    auth_payload = {
        "saved_on": get_today_string(),
        "storage_state": context.storage_state(),
    }
    with open(auth_state_path, "w", encoding="utf-8") as auth_file:
        json.dump(auth_payload, auth_file, indent=2)


def is_auth_redirect(page):
    current_url = (page.url or "").lower()
    return any(host in current_url for host in AUTH_REDIRECT_HOSTS)


def create_browser_context(browser):
    saved_auth_state = load_saved_auth_state(AUTH_STATE_FILE)
    if saved_auth_state:
        print(f"Using saved session from {AUTH_STATE_FILE}")
        return browser.new_context(storage_state=saved_auth_state["storage_state"])

    print("No saved session found. Starting a fresh browser context.")
    return browser.new_context()


def wait_for_portal_ready(page, timeout=15000):
    try:
        page.wait_for_load_state("domcontentloaded", timeout=timeout)
        page.wait_for_load_state("networkidle", timeout=min(timeout, 5000))
    except PlaywrightTimeoutError:
        print("Portal did not become fully idle in time. Continuing anyway.")


def reauthenticate_if_redirected(page, context, portal_name, portal_url):
    if not is_auth_redirect(page):
        return False

    print(
        f"Authentication expired while accessing {portal_name}. "
        "Complete SSO + DUO in the open browser window."
    )
    login_once(page)
    page.goto(
        portal_url,
        wait_until="domcontentloaded",
        timeout=NAVIGATION_TIMEOUT_MS,
    )
    wait_for_portal_ready(page, timeout=30000)

    if is_auth_redirect(page):
        raise RuntimeError(
            f"Authentication is still not complete for {portal_name}. "
            "Finish DUO approval, then rerun the script."
        )

    save_auth_state(context, AUTH_STATE_FILE)
    print(f"Refreshed authenticated session in {AUTH_STATE_FILE}")
    return True


def navigate_to_portal(page, context, portal_name, portal_url, retries=PORTAL_RETRY_COUNT):
    last_error = None
    for attempt in range(1, retries + 1):
        try:
            print(f"Opening {portal_name} ({attempt}/{retries})")
            page.goto(
                portal_url,
                wait_until="domcontentloaded",
                timeout=NAVIGATION_TIMEOUT_MS,
            )
            wait_for_portal_ready(page, timeout=30000)
            reauthenticate_if_redirected(page, context, portal_name, portal_url)
            return True
        except (PlaywrightTimeoutError, PlaywrightError) as exc:
            last_error = exc
            print(f"Navigation failed for {portal_name}: {exc}")
            if attempt < retries:
                print("Retrying portal navigation...")
                page.wait_for_timeout(5000)
        except RuntimeError as exc:
            last_error = exc
            print(f"Authentication failed for {portal_name}: {exc}")
            if attempt < retries:
                print("Retrying portal navigation after authentication failure...")
                page.wait_for_timeout(5000)

    print(f"Skipping {portal_name} after repeated navigation failures: {last_error}")
    return False


def build_search_url(portal_url, query, earliest="-15m", latest="now"):
    query_text = query if query.strip().lower().startswith("search ") else f"search {query}"
    search_params = urlencode(
        {
            "earliest": earliest,
            "latest": latest,
            "q": query_text,
            "display.page.search.mode": "smart",
            "dispatch.sample_ratio": "1",
        },
        quote_via=quote,
    )

    parsed_url = urlsplit(portal_url)
    return urlunsplit(
        (
            parsed_url.scheme,
            parsed_url.netloc,
            parsed_url.path,
            search_params,
            "",
        )
    )


def run_query_from_url(page, context, portal_name, portal_url, query_config):
    search_url = build_search_url(
        portal_url,
        query_config["query"],
        query_config.get("earliest", "-15m"),
        query_config.get("latest", "now"),
    )
    print(f"Running search from URL fallback for {portal_name}")
    page.goto(
        search_url,
        wait_until="domcontentloaded",
        timeout=NAVIGATION_TIMEOUT_MS,
    )
    wait_for_portal_ready(page, timeout=30000)
    reauthenticate_if_redirected(page, context, portal_name, search_url)
    page.wait_for_timeout(SEARCH_RESULTS_WAIT_MS)
    wait_for_portal_ready(page, timeout=10000)


def move_search_cursor_to_end(page, hide_caret=False):
    page.evaluate(
        """
        ({ inputSelector, hideCaret }) => {
            if (hideCaret && !document.querySelector("#itsi-hide-search-caret")) {
                const style = document.createElement("style");
                style.id = "itsi-hide-search-caret";
                style.textContent = `
                    .ace_cursor {
                        opacity: 0 !important;
                    }
                    pre.ace_editor textarea,
                    textarea[placeholder*='enter search here'],
                    input[placeholder*='enter search here'],
                    textarea[aria-label='Search query'],
                    input[aria-label='Search query'],
                    textarea[data-test='search-field'],
                    input[data-test='search-field'],
                    textarea.search-input,
                    input.search-input {
                        caret-color: transparent !important;
                    }
                `;
                document.head.appendChild(style);
            }

            const editorElement = document.querySelector("pre.ace_editor");
            if (editorElement && window.ace) {
                const editor = window.ace.edit(editorElement);
                const session = editor.getSession();
                const lastRow = session.getLength() - 1;
                const lastColumn = session.getLine(lastRow).length;
                editor.focus();
                editor.resize(true);
                editor.moveCursorTo(lastRow, lastColumn);
                editor.clearSelection();
                editor.renderer.updateFull();
                editor.renderer.scrollCursorIntoView();
                return true;
            }

            const searchField = document.querySelector(inputSelector);
            if (searchField) {
                const endPosition = searchField.value.length;
                searchField.focus();
                if (searchField.setSelectionRange) {
                    searchField.setSelectionRange(endPosition, endPosition);
                }
                searchField.scrollLeft = searchField.scrollWidth;
                return true;
            }

            return false;
        }
        """,
        {
            "inputSelector": SEARCH_INPUT_SELECTOR,
            "hideCaret": hide_caret,
        },
    )


def enter_search_query(page, query):
    # ---------- NG1 (Old ACE Editor) ----------
    try:
        ace_editor = page.locator("pre.ace_editor").first
        if ace_editor.is_visible(timeout=3000):
            page.evaluate(
                """
                (query) => {
                    const editor = window.ace.edit(
                        document.querySelector("pre.ace_editor")
                    );

                    editor.focus();
                    editor.setValue(query, 1);
                    editor.clearSelection();
                }
                """,
                query,
            )
            move_search_cursor_to_end(page)

            print("NG1 editor detected")
            page.wait_for_timeout(500)
            return
    except Exception:
        pass

    # ---------- NG2 (New React Search Box) ----------
    try:
        search_box = page.locator(SEARCH_INPUT_SELECTOR).first
        search_box.wait_for(state="visible", timeout=10000)
        search_box.click()
        search_box.fill(query)
        move_search_cursor_to_end(page)

        print("NG2 editor detected")
    except Exception as exc:
        raise RuntimeError(f"Unable to enter search query: {exc}") from exc

    page.wait_for_timeout(500)


def describe_earliest_time(earliest):
    if earliest.startswith("-") and earliest.endswith("d@h"):
        return f"Last {earliest[1:-3]} days"
    if earliest.startswith("-") and earliest.endswith("h@m"):
        return f"Last {earliest[1:-3]} hours"
    if earliest.startswith("-") and earliest.endswith("m"):
        return f"Last {earliest[1:-1]} minutes"
    return earliest


def set_timepicker_to_relative_time(page, earliest):
    preset_selector = f"a[data-earliest='{earliest}'][data-latest='now']"
    time_label = describe_earliest_time(earliest)

    def click_relative_minutes_preset():
        clicked = page.evaluate(
            """
            (presetSelector) => {
                const candidates = Array.from(document.querySelectorAll(presetSelector));
                const target =
                    candidates.find((element) => {
                        const style = window.getComputedStyle(element);
                        return (
                            style.display !== "none" &&
                            style.visibility !== "hidden" &&
                            element.getClientRects().length > 0
                        );
                    }) ||
                    candidates[candidates.length - 1];

                if (!target) {
                    return false;
                }

                target.click();
                return true;
            }
            """,
            preset_selector,
        )
        if not clicked:
            raise RuntimeError(f"{time_label} preset was not found")

    # ---------- NG1 ----------
    try:
        ng1_time = page.locator("div.shared-timerangepicker").first
        if ng1_time.is_visible(timeout=3000):
            ng1_time.click(force=True)
            page.wait_for_timeout(1000)
            click_relative_minutes_preset()

            print(f"NG1 timepicker detected: {time_label}")
            page.wait_for_timeout(1000)
            return
    except Exception:
        pass

    # ---------- NG2 ----------
    try:
        time_picker = page.locator(TIME_PICKER_BUTTON_SELECTOR).first
        time_picker.wait_for(state="visible", timeout=10000)
        time_picker.click(force=True)
        page.wait_for_timeout(1000)
        click_relative_minutes_preset()

        print(f"NG2 timepicker detected: {time_label}")
        page.wait_for_timeout(1000)
    except Exception as exc:
        raise RuntimeError(f"Unable to set timepicker: {exc}") from exc


def get_query_earliest(query_config):
    return query_config.get("earliest", "-15m")


def prepare_page_for_screenshot(page):
    page.evaluate(
        """
        () => {
            const expandControl =
                document.querySelector("td.icon-col .icon-expand") ||
                document.querySelector("td.icon-col a") ||
                document.querySelector("table tbody tr td.icon-col");
            if (expandControl) {
                expandControl.click();
            }
        }
        """
    )
    page.wait_for_timeout(1000)

    page.evaluate(
        """
        () => {
            document.body.style.zoom = "75%";

            const highlightExisting = document.querySelectorAll(".itsi-host-highlight");
            highlightExisting.forEach((node) => {
                const parent = node.parentNode;
                if (!parent) return;
                parent.replaceChild(document.createTextNode(node.textContent || ""), node);
                parent.normalize();
            });

            const walker = document.createTreeWalker(document.body, NodeFilter.SHOW_TEXT);
            const textNodes = [];
            while (walker.nextNode()) {
                const node = walker.currentNode;
                if (node.nodeValue && node.nodeValue.includes("host=")) {
                    textNodes.push(node);
                }
            }

            textNodes.forEach((node) => {
                const value = node.nodeValue;
                if (!value || !node.parentNode) return;
                const fragment = document.createDocumentFragment();
                const parts = value.split(/(host=[^\\s]+)/g);
                parts.forEach((part) => {
                    if (!part) return;
                    if (part.startsWith("host=")) {
                        const span = document.createElement("span");
                        span.className = "itsi-host-highlight";
                        span.textContent = part;
                        span.style.background = "#ffeb3b";
                        span.style.color = "#111111";
                        span.style.fontWeight = "700";
                        span.style.padding = "2px 4px";
                        span.style.borderRadius = "3px";
                        fragment.appendChild(span);
                    } else {
                        fragment.appendChild(document.createTextNode(part));
                    }
                });
                node.parentNode.replaceChild(fragment, node);
            });
        }
        """
    )
    move_search_cursor_to_end(page, hide_caret=True)
    page.wait_for_timeout(500)


def run_search(page):
    # ---------- NG1 ----------
    try:
        ng1_button = page.locator("button[aria-label='Search Button'], a[aria-label='Search Button']").first
        if ng1_button.is_visible(timeout=3000):
            ng1_button.click()

            print("NG1 search button detected")

            page.wait_for_timeout(SEARCH_RESULTS_WAIT_MS)
            wait_for_portal_ready(page, timeout=10000)
            return
    except Exception:
        pass

    # ---------- NG2 ----------
    try:
        ng2_button = page.locator(SEARCH_BUTTON_SELECTOR).first
        ng2_button.wait_for(state="visible", timeout=10000)
        ng2_button.click(force=True)

        print("NG2 search button detected")
    except Exception:
        try:
            page.locator("button:has(svg)").nth(0).click(force=True, timeout=10000)
            print("NG2 search button detected")
        except Exception as exc:
            raise RuntimeError(f"Unable to run search: {exc}") from exc

    page.wait_for_timeout(SEARCH_RESULTS_WAIT_MS)
    wait_for_portal_ready(page, timeout=10000)


def is_no_results_found(page):
    try:
        no_results_message = page.locator(
            "text=No results found. Try expanding the time range."
        ).first
        return no_results_message.is_visible(timeout=2000)
    except Exception:
        return False


def capture_query_result(page, doc, portal_name, query_config):
    prepare_page_for_screenshot(page)

    screenshot_path = os.path.join(
        SCREENSHOT_DIR,
        f"{portal_name.replace(' ', '_')}_{query_config['name'].replace(' ', '_')}.png"
    )

    page.wait_for_timeout(250)
    move_search_cursor_to_end(page, hide_caret=True)
    page.screenshot(path=screenshot_path)

    doc.add_heading(f"{portal_name} - {query_config['name']}", level=1)
    doc.add_paragraph(query_config["query"])
    doc.add_paragraph(f"Query checked for: {describe_earliest_time(query_config.get('earliest', '-15m'))}")
    search_url = page.url
    link_paragraph = doc.add_paragraph("Splunk URL: ")
    add_hyperlink(link_paragraph, search_url, "Open Splunk search")

    if is_no_results_found(page):
        manual_paragraph = doc.add_paragraph(
            "No results found. Please check manually on this URL: "
        )
        add_hyperlink(manual_paragraph, search_url, search_url)
        print(f"No results found. Please check manually on this URL: {search_url}")

    doc.add_picture(screenshot_path, width=Inches(6))
    print(f"Captured screenshot for {portal_name} - {query_config['name']}")


def get_queries_for_portal(portal):
    configured_query_names = portal.get("queries")
    if not configured_query_names:
        configured_query_names = DEFAULT_QUERY_NAMES

    queries_by_key = {normalize_name(query["name"]): query for query in QUERIES}
    query_overrides = {
        normalize_name(query_name): override
        for query_name, override in portal.get("query_overrides", {}).items()
    }
    selected_queries = []
    missing_query_names = []

    for query_name in configured_query_names:
        query = queries_by_key.get(normalize_name(query_name))
        if query:
            query_config = dict(query)
            query_config.update(query_overrides.get(normalize_name(query_name), {}))
            selected_queries.append(query_config)
        else:
            missing_query_names.append(query_name)

    if missing_query_names:
        raise RuntimeError(
            f"{portal['name']} has unknown query name(s): {', '.join(missing_query_names)}"
        )

    selected_query_labels = ", ".join(query["name"] for query in selected_queries)
    print(f"Selected queries for {portal['name']}: {selected_query_labels}")
    return selected_queries


def run_queries_and_capture(page, context, portal, doc):
    portal_name = portal["name"]
    portal_url = portal["url"]
    selected_queries = get_queries_for_portal(portal)

    if not navigate_to_portal(page, context, portal_name, portal_url):
        print(f"Skipping {portal_name}: could not open portal")
        return

    use_url_search = False
    active_earliest = None

    for q in selected_queries:
        try:
            print(f"Running {q['name']} on {portal_name}")

            if use_url_search:
                run_query_from_url(page, context, portal_name, portal_url, q)
            else:
                query_earliest = get_query_earliest(q)
                if active_earliest != query_earliest:
                    try:
                        set_timepicker_to_relative_time(page, query_earliest)
                        active_earliest = query_earliest
                    except RuntimeError as exc:
                        print(f"Timepicker failed on {portal_name}. Using URL search fallback for this portal: {exc}")
                        use_url_search = True
                        run_query_from_url(page, context, portal_name, portal_url, q)
                        active_earliest = None
                        capture_query_result(page, doc, portal_name, q)
                        continue

                enter_search_query(page, q["query"])
                run_search(page)

            capture_query_result(page, doc, portal_name, q)

        except Exception as e:
            print(f"Error in {portal_name} {q['name']}: {e}")


def normalize_name(value):
    return "".join(ch.lower() for ch in value if ch.isalnum())


def prompt_group_alert_earliest(portal_name):
    print(f"\nGroup Alert time picker for {portal_name}")
    print("1 - Default 60 minutes")
    print("2 - Manual Days / Hours / Mins")

    picker_option = input("Choose time picker option: ").strip()
    if not picker_option or picker_option == "1":
        return "-60m"

    if picker_option != "2":
        print("Invalid time picker option. Using default 60 minutes.")
        return "-60m"

    days_value = input("Days: ").strip()
    hours_value = input("Hours: ").strip()
    mins_value = input("Mins: ").strip()

    if not days_value and not hours_value and not mins_value:
        return "-60m"

    try:
        days = int(days_value) if days_value else 0
        hours = int(hours_value) if hours_value else 0
        mins = int(mins_value) if mins_value else 0
    except ValueError as exc:
        raise RuntimeError("Days, Hours, and Mins must be whole numbers.") from exc

    total_minutes = (days * 24 * 60) + (hours * 60) + mins
    if total_minutes <= 0:
        print("No valid Group Alert time entered. Using default 60 minutes.")
        return "-60m"

    if days and not hours and not mins:
        return f"-{days}d@h"
    if hours and not days and not mins:
        return f"-{hours}h@m"
    if mins and not days and not hours:
        return f"-{mins}m"

    return f"-{total_minutes}m"


def select_query_names_for_customer(portal_name):
    print(f"\nSelect queries for {portal_name}:")
    print("1 - All three queries")
    print("2 - Ping Monitor")
    print("3 - Interface Performance")
    print("4 - Device Performance")
    print("5 - Group Alert")
    print("6 - No other query required / end")
    print("Use commas for multiple options, example: 2,4 or 2,5")

    while True:
        raw_value = input("Query options: ").strip()
        if not raw_value:
            print("No option entered. Please enter 1, 2, 3, 4, 5, or 6.")
            continue

        selected_options = [value.strip() for value in raw_value.split(",") if value.strip()]
        invalid_options = [
            option
            for option in selected_options
            if option not in {"1", "2", "3", "4", "5", "6"}
        ]
        if invalid_options:
            print(f"Invalid option(s): {', '.join(invalid_options)}")
            continue

        if "1" in selected_options:
            return DEFAULT_QUERY_NAMES

        selected_query_names = []
        for option in selected_options:
            query_name = QUERY_OPTION_MAP.get(option)
            if query_name and query_name not in selected_query_names:
                selected_query_names.append(query_name)

        if selected_query_names:
            return selected_query_names

        if "6" in selected_options:
            return []

        print("No valid query was selected. Please try again.")


def select_portals():
    print(
        "Enter customer names to check as a comma-separated list "
        "(example: Amgen, Engel, Amex GBT, Allianz ACI, Truist)."
    )
    print("Press Enter without typing anything to run all portals.")
    raw_value = input("Customers: ").strip()
    if not raw_value:
        return PORTALS

    requested_names = [name.strip() for name in raw_value.split(",") if name.strip()]
    matched_portals = []
    missing_names = []

    for requested_name in requested_names:
        requested_key = normalize_name(requested_name)
        match = next(
            (
                portal
                for portal in PORTALS
                if requested_key in normalize_name(portal["name"])
            ),
            None,
        )
        if match:
            matched_portals.append(match)
        else:
            missing_names.append(requested_name)

    if missing_names:
        print(f"No portal match found for: {', '.join(missing_names)}")

    if not matched_portals:
        raise SystemExit("No valid customer names were selected.")

    selected_portals = []
    for portal in matched_portals:
        selected_query_names = select_query_names_for_customer(portal["name"])
        if not selected_query_names:
            print(f"Skipping {portal['name']}: no query selected")
            continue

        portal_with_queries = dict(portal)
        portal_with_queries["queries"] = selected_query_names
        if "Group Alert" in selected_query_names:
            try:
                group_alert_earliest = prompt_group_alert_earliest(portal["name"])
            except RuntimeError as exc:
                print(f"{exc} Using default 60 minutes.")
                group_alert_earliest = "-60m"

            portal_with_queries["query_overrides"] = {
                "Group Alert": {"earliest": group_alert_earliest}
            }
            print(f"Group Alert time for {portal['name']}: {describe_earliest_time(group_alert_earliest)}")

        selected_portals.append(portal_with_queries)

    if not selected_portals:
        raise SystemExit("No customers with query selections were selected.")

    return selected_portals


def main():
    doc = create_document()
    selected_portals = select_portals()
    needs_fresh_login = should_refresh_auth_state(AUTH_STATE_FILE)
    saved_auth_state = load_saved_auth_state(AUTH_STATE_FILE)

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=False)
            if needs_fresh_login:
                if os.path.exists(AUTH_STATE_FILE):
                    auth_saved_on = (
                        saved_auth_state["saved_on"]
                        if saved_auth_state
                        else datetime.fromtimestamp(os.path.getmtime(AUTH_STATE_FILE)).date().isoformat()
                    )
                    print(
                        f"Saved session is from {auth_saved_on}. "
                        "Fresh login required for today."
                    )
                context = browser.new_context()
            else:
                context = create_browser_context(browser)
            page = context.new_page()
            page.set_default_timeout(30000)
            page.set_default_navigation_timeout(NAVIGATION_TIMEOUT_MS)

            # First portal login when no saved session is available.
            if not navigate_to_portal(page, context, selected_portals[0]["name"], selected_portals[0]["url"]):
                raise SystemExit(
                    f"Could not reach the initial portal URL: {selected_portals[0]['url']}. "
                    "Check VPN/network access to the Cisco portal and try again."
                )
            if needs_fresh_login:
                login_once(page)
                save_auth_state(context, AUTH_STATE_FILE)
                print(f"Saved authenticated session to {AUTH_STATE_FILE}")

            for portal in selected_portals:
                print(f"Processing {portal['name']}")
                run_queries_and_capture(page, context, portal, doc)

            browser.close()
    except PermissionError as exc:
        if getattr(exc, "winerror", None) == 5:
            raise SystemExit(
                "Playwright could not start because Windows denied access while "
                "creating its subprocess pipes. This is an environment issue on "
                "this machine, not a portal/query bug. Try running the script from "
                "a normal local folder outside OneDrive with Python 3.12, then "
                "reinstall Playwright in that venv."
            ) from exc
        raise

    # Save Word doc
    report_path = os.path.join(OUTPUT_DIR, "ITSI_Report.docx")
    doc.save(report_path)
    print(f"Report saved at {report_path}")


if __name__ == "__main__":
    main()
